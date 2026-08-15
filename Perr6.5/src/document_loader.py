"""Загрузка и разбиение на чанки документов о ювелирных изделиях."""

from pathlib import Path

from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

DATA_DIR = Path(__file__).resolve().parent.parent / "data" / "Ювелир"

CHUNK_SIZE = 500
CHUNK_OVERLAP = 100


def _load_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _sniff_kind(path: Path) -> str:
    """Определяет реальный тип файла по содержимому, а не по расширению.

    В исходных материалах задания расширения файлов не соответствуют их настоящему формату:
    часть файлов с расширением .pdf на деле являются DOCX (ZIP-контейнер с word/document.xml),
    а один файл с расширением .docx на деле является обычным PDF, другой — обычным текстом.
    Доверять расширению нельзя — определяем тип по первым байтам файла.
    """
    with path.open("rb") as f:
        header = f.read(4)
    if header.startswith(b"%PDF"):
        return "pdf"
    if header.startswith(b"PK\x03\x04"):
        return "docx"
    return "txt"


def load_documents(data_dir: Path = DATA_DIR) -> list[Document]:
    """Загружает все PDF/DOCX/TXT файлы из папки с документами."""
    documents: list[Document] = []
    for path in sorted(data_dir.iterdir()):
        if not path.is_file():
            continue
        kind = _sniff_kind(path)
        if kind == "pdf":
            documents.extend(PyPDFLoader(str(path)).load())
        elif kind == "docx":
            text = _load_docx(path)
            documents.append(Document(page_content=text, metadata={"source": path.name}))
        elif kind == "txt":
            documents.extend(TextLoader(str(path), encoding="utf-8").load())

    # лоадеры пишут в metadata["source"] полный путь к файлу — оставляем только имя,
    # чтобы ссылки на источник в ответах были короткими и одинаковыми для всех типов файлов
    for doc in documents:
        doc.metadata["source"] = Path(doc.metadata.get("source", "")).name

    return documents


def split_documents(documents: list[Document]) -> list[Document]:
    """Разбивает документы на чанки для векторного поиска."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_documents(documents)


def load_and_split(data_dir: Path = DATA_DIR) -> list[Document]:
    return split_documents(load_documents(data_dir))


if __name__ == "__main__":
    chunks = load_and_split()
    print(f"Загружено чанков: {len(chunks)}")
    for chunk in chunks[:3]:
        print("---")
        print(chunk.metadata)
        print(chunk.page_content[:200])
